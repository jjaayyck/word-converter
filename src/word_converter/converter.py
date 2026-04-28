"""Core conversion logic for legacy Word reports."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.table import _Cell
from .config import (
    CELL_CODE_MAPPING,
    FIXED_TEXT_MAPPING,
    HIGH_SCORE_RECOMMENDATION_MAPPING,
    NAME_LABELS,
    SAMPLE_ID_LABELS,
    TABLE_HEADER_MAPPING,
)


@dataclass(slots=True)
class ConversionResult:
    input_path: Path
    output_path: Path
    name: str
    sample_id: str


class WordReportConverter:
    """Convert an old-format Word report to the new format."""

    HEADER_FILL_COLOR = "00B0F0"
    HEADER_FONT_COLOR = "FFFFFF"
    ORANGE_LABEL_FILL_COLOR = "ED7D31"
    GREEN_LABEL_FILL_COLOR = "92D050"
    DARK_BORDER_COLOR = "595959"
    HIGH_SCORE_FONT_COLOR = "00B050"
    RECOMMEND_BORDER_COLOR = "ED7D31"
    EMPHASIS_FONT_COLOR = "ED0000"
    HEADER_BORDER_SIZE_EIGHTHS = 4  # 1/2 pt
    RECOMMEND_BORDER_SIZE_EIGHTHS = 18  # 2 1/4 pt
    FONT_NAME = "微軟正黑體"
    RECOMMENDATION_GREETING_FONT_SIZE_PT = 16
    RECOMMENDATION_INTRO_LINE_SPACING_PT = 16
    HIGH_SCORE_INTRO_BLANK_LINE_SPACING_PT = 19
    DISCLAIMER_FONT_SIZE_PT = 10
    LEFT_LOGO_BASENAME = "威力logo總表"
    RIGHT_LOGO_BASENAME = "心理logo總表"
    RECOMMENDATION_LOGO_BASENAME = "建議logo"
    HIGH_SCORE_ITEMS_PER_PAGE = 2

    LEGACY_MAIN_HEADERS = ["編號", "功能", "細胞解碼位點", "解碼型", "健康優勢評估", "健康優勢評分"]
    NEW_MAIN_HEADERS = ["編號", "心理天賦項目", "細胞解碼位點", "解碼型", "心理潛能優勢評估", "心理潛能優勢評分"]

    ORANGE_INFO_LABELS = {"姓名", "受測者姓名", "受檢者姓名", "出生日期"}
    GREEN_INFO_LABELS = {"送檢編號", "檢體類型"}

    def __init__(
        self,
        table_header_mapping: dict[str, str] | None = None,
        cell_code_mapping: dict[tuple[str, str], str] | None = None,
        fixed_text_mapping: dict[str, str] | None = None,
    ) -> None:
        self.table_header_mapping = table_header_mapping or TABLE_HEADER_MAPPING
        self.normalized_table_header_mapping = {
            self._normalize_label(source): target for source, target in self.table_header_mapping.items()
        }
        self.cell_code_mapping = cell_code_mapping or CELL_CODE_MAPPING
        self.fixed_text_mapping = fixed_text_mapping or FIXED_TEXT_MAPPING
        self.high_score_recommendation_mapping = HIGH_SCORE_RECOMMENDATION_MAPPING
        self.last_cell_code_report: dict[str, Any] = {}

    def convert(self, input_path: str | Path, output_dir: str | Path) -> ConversionResult:
        input_file = Path(input_path)
        if not input_file.exists():
            raise FileNotFoundError(f"Input file not found: {input_file}")

        document = self._load_document(input_file)
        name, sample_id = self._extract_identity(document)
        self._convert_table_headers(document)
        self._convert_cell_codes(document)
        self._apply_fixed_text(document)
        self._replace_recommendation_section(document, name, input_file.parent)
        self._highlight_score_emphasis_text(document)
        self._apply_recommendation_format_overrides(document)
        self._apply_table_styles(document)
        self._apply_first_page_logos(document, input_file.parent)
        self._apply_page_layout(document)
        self._apply_global_font(document)

        output_dir_path = Path(output_dir)
        output_dir_path.mkdir(parents=True, exist_ok=True)
        output_path = output_dir_path / self._build_output_filename(sample_id=sample_id, name=name)
        document.save(str(output_path))

        return ConversionResult(
            input_path=input_file,
            output_path=output_path,
            name=name,
            sample_id=sample_id,
        )

    def preview_output_path(self, input_path: str | Path, output_dir: str | Path) -> Path:
        """Infer converted output path without writing files."""
        input_file = Path(input_path)
        if not input_file.exists():
            raise FileNotFoundError(f"Input file not found: {input_file}")

        document = self._load_document(input_file)
        name, sample_id = self._extract_identity(document)
        output_dir_path = Path(output_dir)
        return output_dir_path / self._build_output_filename(sample_id=sample_id, name=name)

    @staticmethod
    def _load_document(input_file: Path) -> Any:
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            raise ModuleNotFoundError(
                "python-docx is required. Install dependencies with: pip install -r requirements.txt"
            ) from exc

        return Document(str(input_file))

    def _extract_identity(self, document: "DocxDocument") -> tuple[str, str]:
        name: str | None = None
        sample_id: str | None = None

        for table in document.tables:
            t_name, t_id = self._extract_identity_from_table(table)
            name = name or t_name
            sample_id = sample_id or t_id
            if name and sample_id:
                return name, sample_id

        full_text = "\n".join(p.text for p in document.paragraphs if p.text)
        name = name or self._extract_by_labels(full_text, NAME_LABELS)
        sample_id = sample_id or self._extract_by_labels(full_text, SAMPLE_ID_LABELS)

        if name and sample_id:
            return name, sample_id

        missing = []
        if not name:
            missing.append("姓名")
        if not sample_id:
            missing.append("送檢編號")
        raise ValueError(f"無法從文件中擷取欄位：{', '.join(missing)}")

    @staticmethod
    def _normalize_label(value: str) -> str:
        return "".join(value.split())

    @staticmethod
    def _extract_by_labels(text: str, labels: list[str]) -> str | None:
        for label in labels:
            pattern = rf"{re.escape(label)}\s*[:：]\s*([^\n\r\t,，]+)"
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()
        return None

    def _extract_identity_from_table(self, table: "Table") -> tuple[str | None, str | None]:
        name: str | None = None
        sample_id: str | None = None

        normalized_name_labels = {self._normalize_label(label) for label in NAME_LABELS}
        normalized_sample_id_labels = {self._normalize_label(label) for label in SAMPLE_ID_LABELS}

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            for idx, cell_text in enumerate(cells):
                normalized_cell_text = self._normalize_label(cell_text)

                if normalized_cell_text in normalized_name_labels:
                    name = name or self._first_non_empty_to_right(cells, idx)

                if normalized_cell_text in normalized_sample_id_labels:
                    sample_id = sample_id or self._first_non_empty_to_right(cells, idx)

                if name and sample_id:
                    return name, sample_id
        return name, sample_id

    @staticmethod
    def _first_non_empty_to_right(cells: list[str], idx: int) -> str | None:
        for cell_text in cells[idx + 1 :]:
            value = cell_text.strip()
            if value:
                return value
        return None

    def _convert_table_headers(self, document: "DocxDocument") -> None:
        for table in document.tables:
            if not table.rows:
                continue
            for cell in table.rows[0].cells:
                normalized_original = self._normalize_label(cell.text.strip())
                replacement = self.normalized_table_header_mapping.get(normalized_original)
                if replacement:
                    self._replace_cell_text(cell, replacement)

    def _convert_cell_codes(self, document: "DocxDocument") -> None:
        table_reports: list[dict[str, Any]] = []
        main_table_index: int | None = None
        replaced_count = 0
        unmapped_features: list[str] = []
        seen_unmapped: set[str] = set()

        for index, table in enumerate(document.tables):
            header_cells = table.rows[0].cells if table.rows else []
            headers = [cell.text.strip() for cell in header_cells]
            normalized_headers = [self._normalize_label(header) for header in headers]

            is_main_table = self._is_main_table_headers(normalized_headers)
            table_reports.append(
                {
                    "table_index": index,
                    "headers": headers,
                    "is_main_table": is_main_table,
                }
            )

            if not is_main_table or main_table_index is not None:
                continue

            main_table_index = index
            for row in table.rows[1:]:
                if len(row.cells) < 3:
                    continue

                feature_name = row.cells[1].text.strip()
                old_code = row.cells[2].text.strip()
                if not feature_name or not old_code:
                    continue

                new_code = self.cell_code_mapping.get((feature_name, old_code))
                if new_code:
                    self._replace_cell_text(row.cells[2], new_code)
                    self._set_row_height(row, self.GENE_ROW_HEIGHT_EMU)
                    replaced_count += 1
                elif feature_name not in seen_unmapped:
                    seen_unmapped.add(feature_name)
                    unmapped_features.append(feature_name)

        self.last_cell_code_report = {
            "table_count": len(document.tables),
            "tables": table_reports,
            "main_table_index": main_table_index,
            "replaced_count": replaced_count,
            "unmapped_features": unmapped_features,
        }

    def _apply_fixed_text(self, document: "DocxDocument") -> None:
        for paragraph in document.paragraphs:
            replaced = self._replace_fixed_text(paragraph.text)
            if replaced != paragraph.text:
                paragraph.text = replaced

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    replaced = self._replace_fixed_text(cell.text)
                    if replaced != cell.text:
                        self._replace_cell_text(cell, replaced)

    def _replace_fixed_text(self, text: str) -> str:
        replaced = text
        for old, new in self.fixed_text_mapping.items():
            replaced = replaced.replace(old, new)
        return replaced

    def _replace_recommendation_section(
        self,
        document: "DocxDocument",
        name: str,
        input_dir: Path | None = None,
    ) -> None:
        paragraphs = list(getattr(document, "paragraphs", []))
        anchor_index = self._find_disclaimer_anchor_index(paragraphs)
        if anchor_index is None:
            return

        high_features, low_features = self._collect_scored_features(document)
        low_anchor = self._find_low_score_anchor_paragraph(document, low_features)
        if low_anchor is None:
            low_anchor = self._append_paragraph(document, self._build_low_score_intro(low_features))
        high_greeting_anchor = self._find_first_recommendation_greeting_paragraph(document)
        if low_anchor is None:
            return
        if high_greeting_anchor is None:
            self._insert_paragraph_before_anchor(
                document,
                low_anchor,
                self._build_recommendation_greeting(name),
                font_size_pt=self.RECOMMENDATION_GREETING_FONT_SIZE_PT,
            )
            high_greeting_anchor = self._find_first_recommendation_greeting_paragraph(document)
            if high_greeting_anchor is None:
                return

        high_greeting_anchor.text = self._build_recommendation_greeting(name)
        self._style_paragraph_text(high_greeting_anchor, font_size_pt=self.RECOMMENDATION_GREETING_FONT_SIZE_PT)

        self._remove_existing_high_block_between_anchors(document, high_greeting_anchor, low_anchor)

        high_intro_paragraph = None
        for text, font_size_pt in self._build_recommendation_paragraphs(high_features):
            inserted = self._insert_paragraph_before_anchor(document, low_anchor, text, font_size_pt=font_size_pt)
            if inserted is not None and "優勢評估分數較高，在此，也提供給您改善及建議方針：" in text:
                high_intro_paragraph = inserted

        if high_intro_paragraph is not None:
            blank_after_high_intro = self._insert_paragraph_before_anchor(document, low_anchor, "")
            if blank_after_high_intro is not None:
                self._set_paragraph_spacing_pt(blank_after_high_intro, self.HIGH_SCORE_INTRO_BLANK_LINE_SPACING_PT)

        recommendation_logo = self._resolve_logo_path(self.RECOMMENDATION_LOGO_BASENAME, input_dir) if input_dir else None
        self._insert_high_score_tables_before_anchor(document, high_features, low_anchor, recommendation_logo)
        rebuilt_low_anchor = self._insert_low_score_recommendations_before_anchor(
            document,
            low_anchor,
            low_features,
            recommendation_logo,
        )
        if rebuilt_low_anchor is not None:
            low_anchor = rebuilt_low_anchor
        self._insert_paragraph_before_anchor(
            document,
            low_anchor,
            self._build_recommendation_greeting(name),
            font_size_pt=self.RECOMMENDATION_GREETING_FONT_SIZE_PT,
        )
        self._remove_extra_recommendation_greetings(document, keep_count=2)

        if low_anchor is None:
            self._append_page_break(document)

    @staticmethod
    def _append_paragraph(document: "DocxDocument", text: str) -> Any | None:
        if hasattr(document, "add_paragraph"):
            return document.add_paragraph(text)
        if hasattr(document, "paragraphs") and isinstance(document.paragraphs, list):
            paragraph_cls = type(document.paragraphs[0]) if document.paragraphs else None
            if paragraph_cls is not None:
                paragraph = paragraph_cls(text=text)
            else:
                from types import SimpleNamespace

                paragraph = SimpleNamespace(text=text)
            document.paragraphs.append(paragraph)
            return paragraph
        return None

    def _find_first_recommendation_greeting_paragraph(self, document: "DocxDocument") -> Any | None:
        greeting_pattern = re.compile(r"^_+.+_+\s*貴賓您好：\s*$")
        for paragraph in getattr(document, "paragraphs", []):
            if greeting_pattern.match(getattr(paragraph, "text", "").strip()):
                return paragraph
        return None

    def _remove_existing_high_block_between_anchors(
        self,
        document: "DocxDocument",
        start_anchor: Any,
        end_anchor: Any,
    ) -> None:
        paragraphs = getattr(document, "paragraphs", [])
        if start_anchor not in paragraphs or end_anchor not in paragraphs:
            return
        start_idx = paragraphs.index(start_anchor)
        end_idx = paragraphs.index(end_anchor)
        if end_idx <= start_idx + 1:
            return
        if isinstance(paragraphs, list):
            del paragraphs[start_idx + 1 : end_idx]
            return
        for paragraph in list(paragraphs[start_idx + 1 : end_idx]):
            self._remove_paragraph(document, paragraph)

    def _remove_extra_recommendation_greetings(self, document: "DocxDocument", keep_count: int) -> None:
        greeting_pattern = re.compile(r"^_+.+_+\s*貴賓您好：\s*$")
        greetings = [
            paragraph
            for paragraph in list(getattr(document, "paragraphs", []))
            if greeting_pattern.match(getattr(paragraph, "text", "").strip())
        ]
        for paragraph in greetings[keep_count:]:
            self._remove_paragraph(document, paragraph)

    @staticmethod
    def _remove_paragraph(document: "DocxDocument", paragraph: Any) -> None:
        if hasattr(document, "paragraphs") and isinstance(document.paragraphs, list):
            for index, current in enumerate(document.paragraphs):
                if current is paragraph:
                    del document.paragraphs[index]
                    return
            if paragraph in document.paragraphs:
                document.paragraphs.remove(paragraph)
                return
        if hasattr(paragraph, "_element"):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

    def _find_existing_greeting_before_low_anchor(self, document: "DocxDocument", low_anchor: Any | None) -> Any | None:
        if low_anchor is None:
            return None
        paragraphs = list(getattr(document, "paragraphs", []))
        if low_anchor not in paragraphs:
            return None
        low_anchor_index = paragraphs.index(low_anchor)
        greeting_pattern = re.compile(r"^_+.+_+\s*貴賓您好：\s*$")

        for paragraph in reversed(paragraphs[:low_anchor_index]):
            if greeting_pattern.match(getattr(paragraph, "text", "").strip()):
                return paragraph
        return None

    def _remove_existing_high_block_between_disclaimer_and_low_anchor(
        self,
        document: "DocxDocument",
        disclaimer_index: int,
        low_anchor: Any | None,
    ) -> None:
        paragraphs = getattr(document, "paragraphs", [])
        if low_anchor is None or low_anchor not in paragraphs:
            return

        low_anchor_index = paragraphs.index(low_anchor)
        if low_anchor_index <= disclaimer_index + 1:
            return

        if isinstance(paragraphs, list):
            del paragraphs[disclaimer_index + 1 : low_anchor_index]
            return

        for paragraph in list(paragraphs[disclaimer_index + 1 : low_anchor_index]):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

    def _find_disclaimer_anchor_index(self, paragraphs: list[Any]) -> int | None:
        disclaimer_tokens = ("本報告所提供之心理天賦優勢分析", "本報告依細胞分子生物學分析及統計資料")
        for idx, paragraph in enumerate(paragraphs):
            text = getattr(paragraph, "text", "")
            if any(token in text for token in disclaimer_tokens):
                return idx
        return None

    def _collect_scored_features(self, document: "DocxDocument") -> tuple[list[str], list[str]]:
        high_features: list[str] = []
        low_features: list[str] = []

        for table in getattr(document, "tables", []):
            if not table.rows:
                continue
            headers = [self._normalize_label(cell.text.strip()) for cell in table.rows[0].cells]
            if not self._is_main_table_headers(headers):
                continue

            for row in table.rows[1:]:
                if len(row.cells) < 5:
                    continue
                feature = row.cells[1].text.strip()
                score = row.cells[4].text.strip()
                if not feature:
                    continue
                if score == "高":
                    high_features.append(feature)
                elif score == "低":
                    low_features.append(feature)
            break

        return high_features, low_features

    def _find_low_score_anchor_paragraph(self, document: "DocxDocument", low_features: list[str]) -> Any | None:
        anchor_tokens = (
            "優勢評估分數較低，在此，也提供給您改善及建議方針：",
            "健康優勢評估分數較低，在此，也提供給您改善及建議方針：",
        )
        for paragraph in getattr(document, "paragraphs", []):
            text = getattr(paragraph, "text", "")
            if any(token in text for token in anchor_tokens):
                return paragraph

        # fallback: keep compatibility with documents that only contain a short intro marker
        fallback_tokens = ["感謝您接受健康趨勢細胞解碼檢測", "感謝您接受心理潛能細胞解碼檢測"]
        for paragraph in getattr(document, "paragraphs", []):
            text = getattr(paragraph, "text", "")
            if any(token in text for token in fallback_tokens):
                return paragraph

        return None

    @staticmethod
    def _insert_paragraph_before_anchor(
        document: "DocxDocument",
        anchor_paragraph: Any | None,
        text: str,
        font_size_pt: int | None = None,
    ) -> Any | None:
        if anchor_paragraph is not None and hasattr(anchor_paragraph, "insert_paragraph_before"):
            paragraph = anchor_paragraph.insert_paragraph_before(text)
            if font_size_pt is not None:
                WordReportConverter._style_paragraph_text(paragraph, font_size_pt=font_size_pt)
            return paragraph

        if hasattr(document, "paragraphs") and isinstance(document.paragraphs, list):
            paragraph_cls = type(document.paragraphs[0]) if document.paragraphs else None
            if paragraph_cls is not None:
                new_paragraph = paragraph_cls(text=text)
            else:
                from types import SimpleNamespace

                new_paragraph = SimpleNamespace(text=text)

            if anchor_paragraph is not None and anchor_paragraph in document.paragraphs:
                anchor_index = document.paragraphs.index(anchor_paragraph)
                document.paragraphs.insert(anchor_index, new_paragraph)
            else:
                document.paragraphs.append(new_paragraph)
            return new_paragraph

        if hasattr(document, "add_paragraph"):
            paragraph = document.add_paragraph(text)
            if font_size_pt is not None:
                WordReportConverter._style_paragraph_text(paragraph, font_size_pt=font_size_pt)
            return paragraph
        return None

    def _insert_high_score_tables_before_anchor(
        self,
        document: "DocxDocument",
        high_features: list[str],
        anchor_paragraph: Any | None,
        recommendation_logo: Path | None = None,
    ) -> None:
        if not hasattr(document, "add_table"):
            self._insert_page_break_before_anchor(document, anchor_paragraph)
            return

        feature_items = high_features or ["綜合能力"]
        items_on_current_page = 0
        for index, feature in enumerate(feature_items):
            if items_on_current_page == 0:
                logo_anchor_paragraph = self._insert_paragraph_before_anchor(document, anchor_paragraph, "")
                if logo_anchor_paragraph is not None:
                    self._add_recommendation_logo_to_paragraph(logo_anchor_paragraph, recommendation_logo)

            header_table = document.add_table(rows=1, cols=1)
            if anchor_paragraph is not None and hasattr(anchor_paragraph, "_p") and hasattr(header_table, "_tbl"):
                anchor_paragraph._p.addprevious(header_table._tbl)
            header_cell = header_table.rows[0].cells[0]

            self._replace_cell_text(header_cell, feature)
            self._set_cell_fill(header_cell, self.HEADER_FILL_COLOR)
            self._style_cell_text(header_cell, bold=True, font_color=self.HEADER_FONT_COLOR, font_size_pt=18)
            self._set_table_border(
                header_table,
                color=self.DARK_BORDER_COLOR,
                size_eighths=self.HEADER_BORDER_SIZE_EIGHTHS,
            )
            self._set_row_height_pt(header_table.rows[0], 20)
            self._set_cell_paragraph_line_spacing_pt(header_cell, 20)

            spacer = self._insert_or_append_spacer_paragraph(document, anchor_paragraph)
            self._set_paragraph_spacing_pt(spacer, line_spacing_pt=6)

            suggestion_table = document.add_table(rows=1, cols=1)
            if anchor_paragraph is not None and hasattr(anchor_paragraph, "_p") and hasattr(suggestion_table, "_tbl"):
                anchor_paragraph._p.addprevious(suggestion_table._tbl)
            suggestion_cell = suggestion_table.rows[0].cells[0]
            self._replace_cell_text(suggestion_cell, self._build_high_score_suggestion_text(feature))
            self._style_cell_text(suggestion_cell)
            self._set_cell_paragraph_line_spacing_pt(suggestion_cell, 20)
            self._set_table_border(
                suggestion_table,
                color=self.RECOMMEND_BORDER_COLOR,
                size_eighths=self.RECOMMEND_BORDER_SIZE_EIGHTHS,
            )
            self._set_row_height_cm(suggestion_table.rows[0], 7)
            items_on_current_page += 1

            if index < len(feature_items) - 1:
                if items_on_current_page >= self.HIGH_SCORE_ITEMS_PER_PAGE:
                    self._insert_page_break_before_anchor(document, anchor_paragraph)
                    items_on_current_page = 0
                else:
                    between = self._insert_or_append_spacer_paragraph(document, anchor_paragraph)
                    self._set_paragraph_spacing_pt(between, line_spacing_pt=19)
            else:
                self._insert_page_break_before_anchor(document, anchor_paragraph)

    def _add_recommendation_logo_to_paragraph(self, paragraph: Any, recommendation_logo: Path | None) -> None:
        if recommendation_logo is None or not recommendation_logo.exists():
            return
        if not hasattr(paragraph, "add_run"):
            return

        run = paragraph.add_run()
        inline_shape = self._add_inline_picture(run, recommendation_logo, width_cm=18.84, height_cm=4.92)
        self._convert_inline_to_floating_anchor(
            inline_shape,
            x_cm=0,
            y_cm=-4.23,
            horizontal_relative="column",
            vertical_relative="paragraph",
            behind_text=True,
        )

    def _insert_low_score_recommendations_before_anchor(
        self,
        document: "DocxDocument",
        low_anchor: Any | None,
        low_features: list[str],
        recommendation_logo: Path | None,
    ) -> Any | None:
        if low_anchor is None or not hasattr(low_anchor, "_p"):
            return low_anchor
        low_anchor_text = getattr(low_anchor, "text", "")
        if "優勢評估分數較低" not in low_anchor_text and "健康優勢評估分數較低" not in low_anchor_text:
            return low_anchor

        low_section_blocks, insertion_anchor, low_suggestion_mapping = self._collect_low_section_blocks_and_suggestions(
            low_anchor,
            low_features,
        )
        if not low_section_blocks:
            return low_anchor

        for block in low_section_blocks:
            parent = block.getparent()
            if parent is not None:
                parent.remove(block)

        new_low_intro = self._insert_paragraph_before_block(
            document,
            insertion_anchor,
            self._build_low_score_intro(low_features),
        )

        feature_items = low_features or ["綜合能力"]
        for index, feature in enumerate(feature_items):
            if index % self.HIGH_SCORE_ITEMS_PER_PAGE == 0:
                logo_paragraph = self._insert_paragraph_before_block(document, insertion_anchor, "")
                if logo_paragraph is not None:
                    self._add_recommendation_logo_to_paragraph(logo_paragraph, recommendation_logo)

            header_table = self._insert_table_before_block(document, insertion_anchor)
            if header_table is not None:
                header_cell = header_table.rows[0].cells[0]
                self._replace_cell_text(header_cell, feature)
                self._set_cell_fill(header_cell, self.HEADER_FILL_COLOR)
                self._style_cell_text(header_cell, bold=True, font_color=self.HEADER_FONT_COLOR, font_size_pt=18)
                self._set_table_border(
                    header_table,
                    color=self.DARK_BORDER_COLOR,
                    size_eighths=self.HEADER_BORDER_SIZE_EIGHTHS,
                )
                self._set_row_height_pt(header_table.rows[0], 20)
                self._set_cell_paragraph_line_spacing_pt(header_cell, 20)

            spacer = self._insert_paragraph_before_block(document, insertion_anchor, "")
            if spacer is not None:
                self._set_paragraph_spacing_pt(spacer, line_spacing_pt=6)

            suggestion_table = self._insert_table_before_block(document, insertion_anchor)
            if suggestion_table is not None:
                suggestion_cell = suggestion_table.rows[0].cells[0]
                self._replace_cell_text(
                    suggestion_cell,
                    low_suggestion_mapping.get(feature, ""),
                )
                self._style_cell_text(suggestion_cell)
                self._set_cell_paragraph_line_spacing_pt(suggestion_cell, 20)
                self._set_table_border(
                    suggestion_table,
                    color=self.RECOMMEND_BORDER_COLOR,
                    size_eighths=self.RECOMMEND_BORDER_SIZE_EIGHTHS,
                )
                self._set_row_height_cm(suggestion_table.rows[0], 7)

            if index < len(feature_items) - 1:
                if (index + 1) % self.HIGH_SCORE_ITEMS_PER_PAGE == 0:
                    page_break_paragraph = self._insert_paragraph_before_block(document, insertion_anchor, "")
                    if page_break_paragraph is not None and hasattr(page_break_paragraph, "add_run"):
                        run = page_break_paragraph.add_run()
                        try:
                            from docx.enum.text import WD_BREAK

                            run.add_break(WD_BREAK.PAGE)
                        except Exception:
                            run.add_break()
                else:
                    between = self._insert_paragraph_before_block(document, insertion_anchor, "")
                    if between is not None:
                        self._set_paragraph_spacing_pt(between, line_spacing_pt=19)

        return new_low_intro

    def _collect_low_section_blocks_and_suggestions(
        self,
        low_anchor: Any,
        low_features: list[str],
    ) -> tuple[list[Any], Any | None, dict[str, str]]:
        normalized_features = [feature.strip() for feature in low_features if feature.strip()]
        if not normalized_features:
            return [low_anchor._p], low_anchor._p.getnext(), {}

        blocks: list[Any] = []
        mapping: dict[str, str] = {}
        current = low_anchor._p
        feature_index = 0
        waiting_suggestion = False
        end_block = low_anchor._p

        while current is not None:
            blocks.append(current)
            tag_name = getattr(current, "tag", "")
            if tag_name.endswith("}tbl") and feature_index < len(normalized_features):
                table_text = "".join(current.itertext()).strip()
                feature_name = normalized_features[feature_index]
                if not waiting_suggestion and feature_name in table_text:
                    waiting_suggestion = True
                elif waiting_suggestion:
                    mapping[feature_name] = table_text
                    waiting_suggestion = False
                    feature_index += 1
                    if feature_index >= len(normalized_features):
                        end_block = current
                        break
            current = current.getnext()

        if feature_index < len(normalized_features):
            return [low_anchor._p], low_anchor._p.getnext(), mapping

        return blocks, end_block.getnext(), mapping

    @staticmethod
    def _insert_paragraph_before_block(document: "DocxDocument", anchor_block: Any | None, text: str) -> Any | None:
        if not hasattr(document, "add_paragraph"):
            return None
        paragraph = document.add_paragraph(text)
        if anchor_block is not None and hasattr(paragraph, "_p"):
            anchor_block.addprevious(paragraph._p)
        return paragraph

    @staticmethod
    def _insert_table_before_block(document: "DocxDocument", anchor_block: Any | None) -> Any | None:
        if not hasattr(document, "add_table"):
            return None
        table = document.add_table(rows=1, cols=1)
        if anchor_block is not None and hasattr(table, "_tbl"):
            anchor_block.addprevious(table._tbl)
        return table

    def _build_high_score_suggestion_text(self, feature: str) -> str:
        feature_key = feature.strip()
        if feature_key in self.high_score_recommendation_mapping:
            return self.high_score_recommendation_mapping[feature_key]
        return f"◆ {feature_key}表現佳，建議持續練習並規律追蹤，將優勢穩定轉化為日常表現。"

    @staticmethod
    def _insert_page_break_before_anchor(document: "DocxDocument", anchor_paragraph: Any | None) -> None:
        if anchor_paragraph is not None and hasattr(anchor_paragraph, "insert_paragraph_before"):
            paragraph = anchor_paragraph.insert_paragraph_before("")
            if hasattr(paragraph, "add_run"):
                run = paragraph.add_run()
                try:
                    from docx.enum.text import WD_BREAK

                    run.add_break(WD_BREAK.PAGE)
                except Exception:
                    run.add_break()
            return

        if hasattr(document, "paragraphs") and isinstance(document.paragraphs, list):
            from types import SimpleNamespace

            paragraph = SimpleNamespace(text="\f")
            if anchor_paragraph is not None and anchor_paragraph in document.paragraphs:
                anchor_index = document.paragraphs.index(anchor_paragraph)
                document.paragraphs.insert(anchor_index, paragraph)
            else:
                document.paragraphs.append(paragraph)
            return

        WordReportConverter._append_page_break(document)

    @staticmethod
    def _insert_or_append_spacer_paragraph(document: "DocxDocument", anchor_paragraph: Any | None) -> Any:
        if anchor_paragraph is not None and hasattr(anchor_paragraph, "insert_paragraph_before"):
            return anchor_paragraph.insert_paragraph_before("")

        if hasattr(document, "add_paragraph"):
            return document.add_paragraph("")

        if hasattr(document, "paragraphs") and isinstance(document.paragraphs, list):
            paragraph_cls = type(document.paragraphs[0]) if document.paragraphs else None
            if paragraph_cls is not None:
                paragraph = paragraph_cls(text="")
            else:
                from types import SimpleNamespace

                paragraph = SimpleNamespace(text="")
            if anchor_paragraph is not None and anchor_paragraph in document.paragraphs:
                anchor_index = document.paragraphs.index(anchor_paragraph)
                document.paragraphs.insert(anchor_index, paragraph)
            else:
                document.paragraphs.append(paragraph)
            return paragraph

        raise AttributeError("Document does not support paragraph insertion")

    @staticmethod
    def _append_page_break(document: "DocxDocument") -> None:
        if hasattr(document, "add_paragraph"):
            paragraph = document.add_paragraph("")
            run = paragraph.add_run()
            try:
                from docx.enum.text import WD_BREAK
                run.add_break(WD_BREAK.PAGE)
            except Exception:
                run.add_break()

    def _build_recommendation_greeting(self, name: str) -> str:
        return f"_____{name}_____ 貴賓您好："

    def _apply_recommendation_format_overrides(self, document: "DocxDocument") -> None:
        for paragraph in getattr(document, "paragraphs", []):
            text = getattr(paragraph, "text", "")
            if self._is_recommendation_intro_paragraph(text):
                self._set_paragraph_spacing_pt(paragraph, self.RECOMMENDATION_INTRO_LINE_SPACING_PT)
            if self._is_disclaimer_paragraph(text):
                self._style_paragraph_text(paragraph, font_size_pt=self.DISCLAIMER_FONT_SIZE_PT)

    @staticmethod
    def _is_recommendation_intro_paragraph(text: str) -> bool:
        return (
            "優勢評估分數較高，在此，也提供給您改善及建議方針：" in text
            or "優勢評估分數較低，在此，也提供給您改善及建議方針：" in text
        )

    @staticmethod
    def _is_disclaimer_paragraph(text: str) -> bool:
        return "本報告所提供之心理天賦優勢分析" in text

    def _build_recommendation_paragraphs(
        self,
        high_features: list[str],
    ) -> list[tuple[str, int | None]]:
        high_text = "、".join(high_features) if high_features else "綜合能力"
        high_count_text = str(len(high_features)) if high_features else "多"
        return [
            (
                "感謝您接受心理潛能細胞解碼檢測，由檢測結果得知，"
                f"您在此次的分析項目中，{high_text}等共{high_count_text}項優勢評估分數較高，"
                "在此，也提供給您改善及建議方針：",
                None,
            ),
        ]

    @staticmethod
    def _build_low_score_intro(low_features: list[str]) -> str:
        low_text = "、".join(low_features) if low_features else "綜合能力"
        low_count_text = str(len(low_features)) if low_features else "多"
        return (
            "感謝您接受心理潛能細胞解碼檢測，由檢測結果得知，"
            f"您在此次的分析項目中，{low_text}等共{low_count_text}項優勢評估分數較低，"
            "在此，也提供給您改善及建議方針："
        )

    @classmethod
    def _style_paragraph_text(cls, paragraph: Any, font_size_pt: int | None = None) -> None:
        runs = getattr(paragraph, "runs", [])
        if not runs and getattr(paragraph, "text", "") and hasattr(paragraph, "add_run"):
            run = paragraph.add_run(paragraph.text)
            paragraph.text = ""
            runs = [run]
        for run in runs:
            cls._set_run_font(run, font_size_pt=font_size_pt)

    @staticmethod
    def _set_table_border(table: Any, color: str, size_eighths: int) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl_pr = table._tbl.tblPr if hasattr(table, "_tbl") else None
        if tbl_pr is None:
            return

        for old_border in tbl_pr.findall(qn("w:tblBorders")):
            tbl_pr.remove(old_border)

        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge_el = OxmlElement(f"w:{edge}")
            edge_el.set(qn("w:val"), "single")
            edge_el.set(qn("w:sz"), str(size_eighths))
            edge_el.set(qn("w:space"), "0")
            edge_el.set(qn("w:color"), color)
            borders.append(edge_el)
        tbl_pr.append(borders)

    @staticmethod
    def _set_paragraph_spacing_pt(paragraph: Any, line_spacing_pt: int) -> None:
        if not hasattr(paragraph, "paragraph_format"):
            return
        from docx.shared import Pt

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(line_spacing_pt)

    @staticmethod
    def _set_row_height_pt(row: Any, height_pt: int) -> None:
        from docx.shared import Pt

        row.height = Pt(height_pt)
        if hasattr(row, "height_rule"):
            row.height_rule = 2

    @staticmethod
    def _set_row_height_cm(row: Any, height_cm: float) -> None:
        from docx.shared import Cm

        row.height = Cm(height_cm)
        if hasattr(row, "height_rule"):
            row.height_rule = 2

    @staticmethod
    def _set_cell_paragraph_line_spacing_pt(cell: Any, line_spacing_pt: int) -> None:
        for paragraph in getattr(cell, "paragraphs", []):
            WordReportConverter._set_paragraph_spacing_pt(paragraph, line_spacing_pt)

    def _highlight_score_emphasis_text(self, document: "DocxDocument") -> None:
        targets = ("優勢評估分數較高", "優勢評估分數較低")
        for paragraph in getattr(document, "paragraphs", []):
            for target in targets:
                self._highlight_text_in_paragraph(paragraph, target)

        for table in getattr(document, "tables", []):
            for row in getattr(table, "rows", []):
                for cell in getattr(row, "cells", []):
                    for paragraph in getattr(cell, "paragraphs", []):
                        for target in targets:
                            self._highlight_text_in_paragraph(paragraph, target)

    def _highlight_text_in_paragraph(self, paragraph: Any, target: str) -> None:
        text = getattr(paragraph, "text", "")
        if target not in text or not hasattr(paragraph, "add_run"):
            return

        try:
            from docx.shared import RGBColor
        except ModuleNotFoundError:
            return

        if hasattr(paragraph, "clear"):
            paragraph.clear()
        else:
            paragraph.text = ""

        start = 0
        while True:
            idx = text.find(target, start)
            if idx == -1:
                remaining = text[start:]
                if remaining:
                    paragraph.add_run(remaining)
                break
            if idx > start:
                paragraph.add_run(text[start:idx])
            highlight_run = paragraph.add_run(target)
            highlight_run.bold = True
            if getattr(highlight_run, "font", None) is not None:
                highlight_run.font.color.rgb = RGBColor.from_string(self.EMPHASIS_FONT_COLOR)
            start = idx + len(target)

    def _apply_table_styles(self, document: "DocxDocument") -> None:
        for table in document.tables:
            if not table.rows:
                continue

            header_cells = table.rows[0].cells
            normalized_headers = [self._normalize_label(cell.text.strip()) for cell in header_cells]
            if self._is_main_table_headers(normalized_headers):
                self._style_main_table_header_row(header_cells)
                self._style_main_table_score_cells(table)

            self._style_info_label_cells(table)

    def _style_main_table_header_row(self, header_cells: list[Any]) -> None:
        for cell in header_cells:
            self._set_cell_fill(cell, self.HEADER_FILL_COLOR)
            self._style_cell_text(cell, bold=True, font_color=self.HEADER_FONT_COLOR)

    def _style_info_label_cells(self, table: Any) -> None:
        for row in table.rows:
            for cell in row.cells:
                normalized_label = self._normalize_label(cell.text.strip())
                if normalized_label in self.ORANGE_INFO_LABELS:
                    self._set_cell_fill(cell, self.ORANGE_LABEL_FILL_COLOR)
                    self._style_cell_text(cell)
                elif normalized_label in self.GREEN_INFO_LABELS:
                    self._set_cell_fill(cell, self.GREEN_LABEL_FILL_COLOR)
                    self._style_cell_text(cell)

    def _style_main_table_score_cells(self, table: Any) -> None:
        for row in table.rows[1:]:
            if len(row.cells) < 5:
                continue
            score_value = row.cells[4].text.strip()
            if score_value == "高":
                self._style_cell_text(row.cells[4], font_color=self.HIGH_SCORE_FONT_COLOR)

    @classmethod
    def _is_main_table_headers(cls, normalized_headers: list[str]) -> bool:
        if len(normalized_headers) < 6:
            return False
        legacy = cls.LEGACY_MAIN_HEADERS
        new = cls.NEW_MAIN_HEADERS
        return all(
            normalized_headers[col] in {legacy[col], new[col]}
            for col in range(6)
        )

    @classmethod
    def _style_cell_text(
        cls,
        cell: Any,
        bold: bool | None = None,
        font_color: str | None = None,
        font_size_pt: int | None = None,
    ) -> None:
        paragraphs = getattr(cell, "paragraphs", None)
        if not paragraphs:
            return

        for paragraph in paragraphs:
            runs = getattr(paragraph, "runs", [])
            if not runs and paragraph.text:
                run = paragraph.add_run(paragraph.text)
                paragraph.text = ""
                runs = [run]
            for run in runs:
                cls._set_run_font(run, bold=bold, font_color=font_color, font_size_pt=font_size_pt)

    @classmethod
    def _set_run_font(
        cls,
        run: Any,
        bold: bool | None = None,
        font_color: str | None = None,
        font_size_pt: int | None = None,
    ) -> None:
        font = getattr(run, "font", None)
        if font is None:
            return

        if bold is not None:
            font.bold = bold
        font.name = cls.FONT_NAME
        if font_size_pt is not None:
            from docx.shared import Pt

            font.size = Pt(font_size_pt)

        if font_color:
            from docx.shared import RGBColor

            font.color.rgb = RGBColor.from_string(font_color)

        r_pr = getattr(run, "_element", None)
        if r_pr is None:
            return

        r_fonts = run._element.rPr.rFonts if run._element.rPr is not None else None
        if r_fonts is None:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            if run._element.rPr is None:
                run._element.get_or_add_rPr()
            r_fonts = OxmlElement("w:rFonts")
            run._element.rPr.append(r_fonts)
            r_fonts.set(qn("w:ascii"), cls.FONT_NAME)
            r_fonts.set(qn("w:hAnsi"), cls.FONT_NAME)
            r_fonts.set(qn("w:eastAsia"), cls.FONT_NAME)
            r_fonts.set(qn("w:cs"), cls.FONT_NAME)
        else:
            from docx.oxml.ns import qn

            r_fonts.set(qn("w:ascii"), cls.FONT_NAME)
            r_fonts.set(qn("w:hAnsi"), cls.FONT_NAME)
            r_fonts.set(qn("w:eastAsia"), cls.FONT_NAME)
            r_fonts.set(qn("w:cs"), cls.FONT_NAME)

    @staticmethod
    def _set_cell_fill(cell: Any, fill_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr() if hasattr(cell, "_tc") else None
        if tc_pr is None:
            return

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for child in tc_pr.findall(qn("w:shd")):
            tc_pr.remove(child)

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr.append(shd)

    def _apply_global_font(self, document: "DocxDocument") -> None:
        for paragraph in document.paragraphs:
            for run in getattr(paragraph, "runs", []):
                self._set_run_font(run)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in getattr(cell, "paragraphs", []):
                        for run in getattr(paragraph, "runs", []):
                            self._set_run_font(run)

    def _has_disclaimer_text(self, document: "DocxDocument") -> bool:
        disclaimer_tokens = ["本報告所提供之心理天賦優勢分析", "本報告依細胞分子生物學分析及統計資料"]
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if any(token in text for token in disclaimer_tokens):
                return True
        return False

    def _apply_first_page_logos(self, document: "DocxDocument", input_dir: Path) -> None:
        left_logo = self._resolve_logo_path(self.LEFT_LOGO_BASENAME, input_dir)
        right_logo = self._resolve_logo_path(self.RIGHT_LOGO_BASENAME, input_dir)
        if left_logo is None or right_logo is None:
            missing = []
            if left_logo is None:
                missing.append(self.LEFT_LOGO_BASENAME)
            if right_logo is None:
                missing.append(self.RIGHT_LOGO_BASENAME)
            raise FileNotFoundError(f"找不到 logo 圖檔：{', '.join(missing)}")

        paragraph = self._find_first_body_logo_paragraph(document)
        if paragraph is None:
            raise ValueError("找不到正文第一個含圖片（w:drawing 或 w:pict）的段落，無法替換 logo。")
        self._remove_paragraph_drawings(paragraph)
        paragraph.text = ""

        left_run = paragraph.add_run()
        self._add_inline_picture(left_run, left_logo, width_cm=9.08, height_cm=2.48)

        spacer = paragraph.add_run(" ")
        spacer.font.size = None

        right_run = paragraph.add_run()
        inline_shape = self._add_inline_picture(right_run, right_logo, width_cm=10.24, height_cm=2.79)
        anchor = self._convert_inline_to_floating_anchor(
            inline_shape,
            x_cm=8.76,
            y_cm=0,
            horizontal_relative="page",
            vertical_relative="paragraph",
            behind_text=False,
        )
        if anchor is None:
            raise RuntimeError("右側 logo 轉換為 floating anchor 失敗。")

    @staticmethod
    def _resolve_logo_path(basename: str, input_dir: Path) -> Path | None:
        exts = (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff")
        search_dirs = [
            input_dir,
            Path.cwd(),
            Path.cwd() / "assets",
            Path.cwd() / "samples",
            Path.cwd() / "samples" / "assets",
        ]
        for folder in search_dirs:
            for ext in exts:
                candidate = folder / f"{basename}{ext}"
                if candidate.exists():
                    return candidate
        return None

    @staticmethod
    def _find_first_body_logo_paragraph(document: "DocxDocument") -> Any | None:
        for paragraph in getattr(document, "paragraphs", []):
            if not hasattr(paragraph, "_p"):
                continue
            if paragraph._p.xpath(".//w:drawing | .//w:pict"):
                return paragraph
        return None

    @staticmethod
    def _remove_paragraph_drawings(paragraph: Any) -> None:
        if not hasattr(paragraph, "_p"):
            return
        p_elem = paragraph._p
        for node in p_elem.xpath(".//w:drawing | .//w:pict"):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)

    @staticmethod
    def _add_inline_picture(run: Any, image_path: Path, width_cm: float, height_cm: float) -> Any:
        from docx.shared import Cm

        picture = run.add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
        if picture is None:
            raise RuntimeError(f"插入圖片失敗：{image_path}")
        return picture

    @staticmethod
    def _convert_inline_to_floating_anchor(
        inline_shape: Any,
        x_cm: float,
        y_cm: float,
        horizontal_relative: str,
        vertical_relative: str,
        behind_text: bool,
    ) -> Any:
        from copy import deepcopy

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm

        inline = inline_shape._inline
        graphic = inline.find(qn("a:graphic"))
        extent = inline.find(qn("wp:extent"))
        doc_pr = inline.find(qn("wp:docPr"))
        c_nv_graphic_frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))

        if graphic is None or extent is None:
            raise ValueError("inline 圖片 XML 缺少必要節點：a:graphic / wp:extent")
        if doc_pr is None:
            doc_pr = OxmlElement("wp:docPr")
            doc_pr.set("id", "1")
            doc_pr.set("name", "Picture 1")
        if c_nv_graphic_frame_pr is None:
            c_nv_graphic_frame_pr = OxmlElement("wp:cNvGraphicFramePr")

        anchor = OxmlElement("wp:anchor")
        anchor.set("simplePos", "0")
        anchor.set("relativeHeight", "251658240")
        anchor.set("behindDoc", "1" if behind_text else "0")
        anchor.set("locked", "0")
        anchor.set("layoutInCell", "1")
        anchor.set("allowOverlap", "1")
        anchor.set("distT", "0")
        anchor.set("distB", "0")
        anchor.set("distL", "0")
        anchor.set("distR", "0")

        simple_pos = OxmlElement("wp:simplePos")
        simple_pos.set("x", "0")
        simple_pos.set("y", "0")
        anchor.append(simple_pos)

        position_h = OxmlElement("wp:positionH")
        position_h.set("relativeFrom", horizontal_relative)
        pos_h_offset = OxmlElement("wp:posOffset")
        pos_h_offset.text = str(int(Cm(x_cm)))
        position_h.append(pos_h_offset)
        anchor.append(position_h)

        position_v = OxmlElement("wp:positionV")
        position_v.set("relativeFrom", vertical_relative)
        pos_v_offset = OxmlElement("wp:posOffset")
        pos_v_offset.text = str(int(Cm(y_cm)))
        position_v.append(pos_v_offset)
        anchor.append(position_v)

        extent_elem = OxmlElement("wp:extent")
        extent_elem.set("cx", extent.get("cx", "0"))
        extent_elem.set("cy", extent.get("cy", "0"))
        anchor.append(extent_elem)

        effect_extent = OxmlElement("wp:effectExtent")
        effect_extent.set("l", "0")
        effect_extent.set("t", "0")
        effect_extent.set("r", "0")
        effect_extent.set("b", "0")
        anchor.append(effect_extent)

        wrap_none = OxmlElement("wp:wrapNone")
        anchor.append(wrap_none)

        anchor.append(deepcopy(doc_pr))
        anchor.append(deepcopy(c_nv_graphic_frame_pr))
        anchor.append(deepcopy(graphic))

        inline.getparent().replace(inline, anchor)
        return anchor

    def _apply_page_layout(self, document: "DocxDocument") -> None:
        sections = list(getattr(document, "sections", []))
        if not sections:
            return

        target_sections = sections
        if self._has_disclaimer_text(document) and len(sections) > 1:
            target_sections = sections[:-1]

        for section in target_sections:
            section.top_margin = self.TOP_MARGIN_EMU
            section.left_margin = self.SIDE_BOTTOM_MARGIN_EMU
            section.right_margin = self.SIDE_BOTTOM_MARGIN_EMU
            section.bottom_margin = self.SIDE_BOTTOM_MARGIN_EMU

    @staticmethod
    def _set_row_height(row: Any, height_emu: int) -> None:
        row.height = height_emu
        if hasattr(row, "height_rule"):
            row.height_rule = 2  # WD_ROW_HEIGHT_RULE.EXACTLY

    @staticmethod
    def _replace_cell_text(cell: "_Cell", text: str) -> None:
        if hasattr(cell, "paragraphs") and cell.paragraphs:
            cell.paragraphs[0].text = text
            for paragraph in cell.paragraphs[1:]:
                paragraph.text = ""
        else:
            cell.text = text

    @staticmethod
    def _sanitize_filename_part(value: str) -> str:
        return re.sub(r'[\\/:*?"<>|\s]+', "", value)

    def _build_output_filename(self, sample_id: str, name: str) -> str:
        sid = self._sanitize_filename_part(sample_id)
        person = self._sanitize_filename_part(name)
        return f"台-{sid}_{person}-天賦30項.docx"

    GENE_ROW_HEIGHT_EMU = int(1.9 * 360000)
    TOP_MARGIN_EMU = int(0.75 * 360000)
    SIDE_BOTTOM_MARGIN_EMU = int(1.0 * 360000)
