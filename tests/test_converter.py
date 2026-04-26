from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import pytest

from src.word_converter.converter import WordReportConverter


@dataclass
class FakeCell:
    text: str


@dataclass
class FakeRow:
    cells: list[FakeCell]
    height: int | None = None
    height_rule: int | None = None


@dataclass
class FakeTable:
    rows: list[FakeRow]


@dataclass
class FakeParagraph:
    text: str


@dataclass
class FakeDocument:
    paragraphs: list[FakeParagraph]
    tables: list[FakeTable]
    sections: list[object] = field(default_factory=list)


def _build_main_table() -> FakeTable:
    return FakeTable(
        rows=[
            FakeRow(
                cells=[
                    FakeCell("編 號"),
                    FakeCell("功 能"),
                    FakeCell("細胞解碼位點"),
                    FakeCell("解碼型"),
                    FakeCell("健康優勢評估"),
                    FakeCell("健康優勢評分"),
                ]
            ),
            FakeRow(cells=[FakeCell("1"), FakeCell("運動神經"), FakeCell("CNTF"), FakeCell("AA"), FakeCell("-"), FakeCell("90")]),
            FakeRow(cells=[FakeCell("2"), FakeCell("反應速度"), FakeCell("CNTF"), FakeCell("AB"), FakeCell("-"), FakeCell("91")]),
            FakeRow(cells=[FakeCell("3"), FakeCell("專注穩定性"), FakeCell("HTR2C"), FakeCell("BB"), FakeCell("-"), FakeCell("92")]),
            FakeRow(cells=[FakeCell("4"), FakeCell("專注力"), FakeCell("HTR2C"), FakeCell("BC"), FakeCell("-"), FakeCell("93")]),
            FakeRow(cells=[FakeCell("5"), FakeCell("協調性"), FakeCell("α-actinin"), FakeCell("CC"), FakeCell("-"), FakeCell("94")]),
            FakeRow(cells=[FakeCell("6"), FakeCell("肢體靈活性"), FakeCell("α-actinin"), FakeCell("CD"), FakeCell("-"), FakeCell("95")]),
        ]
    )


def _build_realistic_table() -> FakeTable:
    return FakeTable(
        rows=[
            FakeRow(
                cells=[
                    FakeCell("姓   名"),
                    FakeCell("王曉明"),
                    FakeCell("性別"),
                    FakeCell("男"),
                    FakeCell("出生日期"),
                    FakeCell("1995-07-15"),
                ]
            ),
            FakeRow(
                cells=[
                    FakeCell("送檢編號"),
                    FakeCell("APT-01-00XXXX"),
                    FakeCell("APT-01-00XXXX"),
                    FakeCell("APT-01-00XXXX"),
                    FakeCell("檢體類型"),
                    FakeCell("口腔黏膜"),
                ]
            ),
        ]
    )


def test_extract_name_from_table_with_spaced_label() -> None:
    converter = WordReportConverter()
    doc = FakeDocument(paragraphs=[], tables=[_build_realistic_table()])

    name, _ = converter._extract_identity(doc)

    assert name == "王曉明"


def test_extract_sample_id_from_table() -> None:
    converter = WordReportConverter()
    doc = FakeDocument(paragraphs=[], tables=[_build_realistic_table()])

    _, sample_id = converter._extract_identity(doc)

    assert sample_id == "APT-01-00XXXX"


def test_output_filename_format() -> None:
    converter = WordReportConverter()

    output = converter._build_output_filename(sample_id="APT-01-00XXXX", name="王曉明")

    assert output == "台-APT-01-00XXXX_王曉明-天賦30項.docx"


def test_convert_table_headers_replaces_main_table_titles() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    doc = FakeDocument(paragraphs=[], tables=[main_table])

    converter._convert_table_headers(doc)

    headers = [cell.text for cell in main_table.rows[0].cells]
    assert headers == ["編 號", "心理天賦項目", "細胞解碼位點", "解碼型", "心理潛能優勢評估", "心理潛能優勢評分"]




def test_convert_table_headers_keeps_name_label_unchanged() -> None:
    converter = WordReportConverter()
    table = FakeTable(rows=[FakeRow(cells=[FakeCell("姓名"), FakeCell("送檢編號")])])
    doc = FakeDocument(paragraphs=[], tables=[table])

    converter._convert_table_headers(doc)

    assert table.rows[0].cells[0].text == "姓名"

def test_convert_cell_codes_only_on_main_table() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    non_main_table = FakeTable(
        rows=[
            FakeRow(cells=[FakeCell("說明"), FakeCell("CNTF")]),
            FakeRow(cells=[FakeCell("備註"), FakeCell("HTR2C")]),
        ]
    )
    doc = FakeDocument(paragraphs=[FakeParagraph("段落中 CNTF 不應被替換")], tables=[non_main_table, main_table])

    converter._convert_cell_codes(doc)

    assert main_table.rows[1].cells[2].text == "MN001"
    assert main_table.rows[2].cells[2].text == "RTS001"
    assert non_main_table.rows[0].cells[1].text == "CNTF"
    assert doc.paragraphs[0].text == "段落中 CNTF 不應被替換"
    assert converter.last_cell_code_report["table_count"] == 2
    assert converter.last_cell_code_report["main_table_index"] == 1


def test_convert_cell_codes_still_works_after_header_replacement() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    doc = FakeDocument(paragraphs=[], tables=[main_table])

    converter._convert_table_headers(doc)
    converter._convert_cell_codes(doc)

    assert main_table.rows[1].cells[2].text == "MN001"
    assert converter.last_cell_code_report["main_table_index"] == 0


def test_convert_cell_codes_distinguishes_duplicate_legacy_codes() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    doc = FakeDocument(paragraphs=[], tables=[main_table])

    converter._convert_cell_codes(doc)

    assert main_table.rows[1].cells[2].text == "MN001"  # 運動神經 + CNTF
    assert main_table.rows[2].cells[2].text == "RTS001"  # 反應速度 + CNTF
    assert main_table.rows[3].cells[2].text == "SF001"  # 專注穩定性 + HTR2C
    assert main_table.rows[4].cells[2].text == "CCT001"  # 專注力 + HTR2C
    assert main_table.rows[5].cells[2].text == "CD001"  # 協調性 + α-actinin
    assert main_table.rows[6].cells[2].text == "PFB001"  # 肢體靈活性 + α-actinin
    assert converter.last_cell_code_report["replaced_count"] == 6
    assert converter.last_cell_code_report["unmapped_features"] == []


def test_convert_cell_codes_collects_unmapped_features() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows.append(
        FakeRow(cells=[FakeCell("7"), FakeCell("不存在功能"), FakeCell("CNTF"), FakeCell("ZZ"), FakeCell("-"), FakeCell("80")])
    )
    doc = FakeDocument(paragraphs=[], tables=[main_table])

    converter._convert_cell_codes(doc)

    assert converter.last_cell_code_report["main_table_index"] == 0
    assert converter.last_cell_code_report["replaced_count"] == 6
    assert converter.last_cell_code_report["unmapped_features"] == ["不存在功能"]


def test_apply_fixed_text_replaces_declaration_and_score_item_texts() -> None:
    converter = WordReportConverter()
    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告僅供參考"),
            FakeParagraph("高分項目代表先天優勢"),
            FakeParagraph("低分項目代表先天不足"),
        ],
        tables=[],
    )

    converter._apply_fixed_text(doc)

    assert doc.paragraphs[0].text == "本報告為天賦 30 項分析結果，僅供健康管理參考。"
    assert doc.paragraphs[1].text == "高分項目代表相對優勢，建議持續強化並轉化為日常表現。"
    assert doc.paragraphs[2].text == "低分項目代表目前較需補強，建議透過訓練與習慣養成逐步改善。"


def test_apply_fixed_text_replaces_legacy_low_score_recommendation_template() -> None:
    converter = WordReportConverter()
    old_text = (
        "感謝您接受健康趨勢細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，"
        "想像力、情感分享力、挫折耐受力、危機處理力、膽量、挑戰力等共六項健康優勢評估分數較低，"
        "在此，也提供給您改善及建議方針："
    )
    doc = FakeDocument(paragraphs=[FakeParagraph(old_text)], tables=[])

    converter._apply_fixed_text(doc)

    assert doc.paragraphs[0].text == (
        "感謝您接受心理潛能細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，"
        "想像力、情感分享力、挫折耐受力、危機處理力、膽量、挑戰力等共六項優勢評估分數較低，"
        "在此，也提供給您改善及建議方針："
    )


def test_apply_fixed_text_replaces_text_inside_table_cells() -> None:
    converter = WordReportConverter()
    table = FakeTable(
        rows=[
            FakeRow(cells=[FakeCell("說明"), FakeCell("如有疑問請洽客服")]),
            FakeRow(cells=[FakeCell("提醒"), FakeCell("高分項目代表先天優勢")]),
        ]
    )
    doc = FakeDocument(paragraphs=[], tables=[table])

    converter._apply_fixed_text(doc)

    assert table.rows[0].cells[1].text == "如需進一步解讀，請聯繫專屬顧問或客服中心。"
    assert table.rows[1].cells[1].text == "高分項目代表相對優勢，建議持續強化並轉化為日常表現。"


def test_replace_recommendation_section_updates_name_and_templates() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows[1].cells[1].text = "高特質"
    main_table.rows[1].cells[4].text = "高"
    main_table.rows[2].cells[1].text = "低特質"
    main_table.rows[2].cells[4].text = "低"
    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph("吳峻維"),
            FakeParagraph("舊版健康管理文案"),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "王曉明")

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "_____王曉明_____ 貴賓您好：" in all_text
    assert "_____\n王曉明\n_____" not in all_text
    assert "感謝您接受心理潛能細胞解碼檢測" in all_text
    assert "王曉明" in all_text
    assert "高特質等共1項優勢評估分數較高" in all_text


def test_replace_recommendation_section_inserts_greeting_before_low_score_section() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows[1].cells[1].text = "高特質"
    main_table.rows[1].cells[4].text = "高"
    main_table.rows[2].cells[1].text = "低特質"
    main_table.rows[2].cells[4].text = "低"
    low_intro = (
        "感謝您接受心理潛能細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，"
        "低特質等共1項優勢評估分數較低，在此，也提供給您改善及建議方針："
    )
    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph(low_intro),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "張西西")

    all_text = [p.text for p in doc.paragraphs]
    assert "_____張西西_____ 貴賓您好：" in all_text
    low_index = all_text.index(low_intro)
    greeting_indices = [idx for idx, text in enumerate(all_text) if text == "_____張西西_____ 貴賓您好："]
    assert greeting_indices
    assert greeting_indices[-1] < low_index


def test_replace_recommendation_section_inserts_greeting_before_legacy_low_score_paragraph() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows[1].cells[4].text = "高"
    main_table.rows[2].cells[4].text = "低"
    legacy_low_intro = (
        "感謝您接受健康趨勢細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，"
        "學習能力、想像力、空間感、肌耐力、優質睡眠、挫折耐受力、危機處理力、膽量、挑戰力等共九項健康優勢評估分數較低，"
        "在此，也提供給您改善及建議方針："
    )
    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph("想像力"),
            FakeParagraph(legacy_low_intro),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "張西西")

    texts = [p.text for p in doc.paragraphs]
    low_idx = texts.index(legacy_low_intro)
    greeting_idx = max(i for i, text in enumerate(texts) if text == "_____張西西_____ 貴賓您好：")
    assert greeting_idx < low_idx


def test_replace_recommendation_section_keeps_template_low_greeting_and_adds_only_one_high_greeting() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows[1].cells[1].text = "爆發力"
    main_table.rows[1].cells[4].text = "高"
    main_table.rows[2].cells[1].text = "低特質"
    main_table.rows[2].cells[4].text = "低"
    low_intro = "感謝您接受心理潛能細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，低特質等共1項優勢評估分數較低，在此，也提供給您改善及建議方針："
    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph("_____舊名字_____ 貴賓您好："),
            FakeParagraph(low_intro),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "張西西")

    texts = [p.text for p in doc.paragraphs]
    greeting = "_____張西西_____ 貴賓您好："
    greeting_indexes = [idx for idx, text in enumerate(texts) if text == greeting]
    assert len(greeting_indexes) == 2

    high_intro_index = next(idx for idx, text in enumerate(texts) if "優勢評估分數較高" in text)
    low_intro_index = next(idx for idx, text in enumerate(texts) if "優勢評估分數較低" in text)
    assert greeting_indexes[0] < high_intro_index
    assert greeting_indexes[1] < low_intro_index
    assert all(greeting not in text for text in texts[high_intro_index + 1 : greeting_indexes[1]])


def test_replace_recommendation_section_separates_high_and_low_insertion_paths() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows[1].cells[1].text = "爆發力"
    main_table.rows[1].cells[4].text = "高"
    main_table.rows[2].cells[1].text = "低特質"
    main_table.rows[2].cells[4].text = "低"
    low_intro = "感謝您接受心理潛能細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，低特質等共1項優勢評估分數較低，在此，也提供給您改善及建議方針："
    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph("_____舊名字_____ 貴賓您好："),
            FakeParagraph(low_intro),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "張西西")

    texts = [p.text for p in doc.paragraphs]
    greeting = "_____張西西_____ 貴賓您好："
    greeting_indexes = [idx for idx, text in enumerate(texts) if text == greeting]
    assert len(greeting_indexes) == 2

    high_intro_index = next(idx for idx, text in enumerate(texts) if "優勢評估分數較高" in text)
    low_intro_index = texts.index(low_intro)
    assert greeting_indexes[0] + 1 == high_intro_index
    assert any(text == "\f" for text in texts[high_intro_index + 1 : greeting_indexes[1]])
    assert greeting_indexes[1] < low_intro_index


def test_apply_fixed_text_replaces_long_declaration_text() -> None:
    converter = WordReportConverter()
    old_text = (
        "o\t本報告依細胞分子生物學分析及統計資料，以口腔黏膜樣本進行檢測，僅供本次健康管理參考，"
        "無臨床診斷效力，亦不可作為醫療診斷依據。如有健康疑慮，請諮詢專業醫師。"
    )
    doc = FakeDocument(paragraphs=[FakeParagraph(old_text)], tables=[])

    converter._apply_fixed_text(doc)

    assert "本報告所提供之心理天賦優勢分析" in doc.paragraphs[0].text
    assert doc.paragraphs[0].text.startswith("o\t本報告所提供之心理天賦優勢分析")


@dataclass
class FakeSection:
    top_margin: int | None = None
    left_margin: int | None = None
    right_margin: int | None = None
    bottom_margin: int | None = None


def test_convert_cell_codes_updates_gene_row_height() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    doc = FakeDocument(paragraphs=[], tables=[main_table])

    converter._convert_cell_codes(doc)

    assert main_table.rows[1].height == int(1.9 * 360000)
    assert main_table.rows[1].height_rule == 2




def test_apply_page_layout_skips_last_section_when_disclaimer_exists() -> None:
    converter = WordReportConverter()
    section_a = FakeSection()
    section_b = FakeSection()
    doc = FakeDocument(
        paragraphs=[FakeParagraph("本報告所提供之心理天賦優勢分析")],
        tables=[],
        sections=[section_a, section_b],
    )

    converter._apply_page_layout(doc)

    assert section_a.top_margin == int(0.75 * 360000)
    assert section_b.top_margin is None

def test_apply_page_layout_updates_margins() -> None:
    converter = WordReportConverter()
    section = FakeSection()
    doc = FakeDocument(paragraphs=[], tables=[], sections=[section])

    converter._apply_page_layout(doc)

    assert section.top_margin == int(0.75 * 360000)
    assert section.left_margin == int(1.0 * 360000)
    assert section.right_margin == int(1.0 * 360000)
    assert section.bottom_margin == int(1.0 * 360000)


def test_replace_recommendation_section_inserts_high_block_before_low_anchor() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    for row in main_table.rows[1:4]:
        row.cells[4].text = "高"
    main_table.rows[4].cells[4].text = "低"

    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph("感謝您接受健康趨勢細胞解碼檢測"),
            FakeParagraph("想像力"),
            FakeParagraph("低分項目代表先天不足"),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "王曉明")

    texts = [p.text for p in doc.paragraphs]
    high_idx = next(i for i, t in enumerate(texts) if "_____王曉明_____ 貴賓您好：" in t)
    low_idx = texts.index("感謝您接受健康趨勢細胞解碼檢測")
    assert high_idx < low_idx


def test_replace_recommendation_section_does_not_append_high_block_to_document_end_when_low_anchor_exists() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows[1].cells[4].text = "高"
    main_table.rows[2].cells[4].text = "低"

    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph("感謝您接受健康趨勢細胞解碼檢測"),
            FakeParagraph("結尾段落"),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "王曉明")

    texts = [p.text for p in doc.paragraphs]
    assert texts[-1] == "結尾段落"
    assert any("_____王曉明_____ 貴賓您好：" in text for text in texts[:-1])


def test_replace_recommendation_section_uses_actual_high_feature_count_text() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    for row in main_table.rows[1:]:
        row.cells[4].text = "高"

    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph("感謝您接受健康趨勢細胞解碼檢測"),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "王曉明")

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "共6項" in all_text


def test_replace_recommendation_section_inserts_page_break_before_low_anchor() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows[1].cells[4].text = "高"
    main_table.rows[2].cells[4].text = "低"

    low_anchor = FakeParagraph("感謝您接受健康趨勢細胞解碼檢測")
    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            low_anchor,
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "王曉明")

    low_idx = doc.paragraphs.index(low_anchor)
    assert low_idx > 1
    assert doc.paragraphs[low_idx - 2].text == "\f"
    assert doc.paragraphs[low_idx - 1].text == "_____王曉明_____ 貴賓您好："


def test_recommendation_section_applies_summary_emphasis_and_non_placeholder_suggestion() -> None:
    from docx import Document

    converter = WordReportConverter()
    doc = Document()
    doc.add_paragraph("本報告所提供之心理天賦優勢分析")
    doc.add_paragraph("感謝您接受心理潛能細胞解碼檢測，以下為低分建議區塊。")
    doc.add_paragraph("文件結尾段落")

    table = doc.add_table(rows=3, cols=6)
    headers = ["編號", "功能", "細胞解碼位點", "解碼型", "健康優勢評估", "健康優勢評分"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    table.rows[1].cells[1].text = "空間感"
    table.rows[1].cells[4].text = "高"
    table.rows[2].cells[1].text = "想像力"
    table.rows[2].cells[4].text = "低"

    converter._replace_recommendation_section(doc, "王曉明")
    converter._highlight_score_emphasis_text(doc)

    texts = [p.text for p in doc.paragraphs]
    summary_idx = next(i for i, text in enumerate(texts) if "優勢評估分數較高" in text)
    low_anchor_idx = next(i for i, text in enumerate(texts) if "以下為低分建議區塊" in text)
    assert summary_idx < low_anchor_idx
    assert texts[-1] == "文件結尾段落"
    assert any("以下為低分建議區塊" in text for text in texts)

    summary_paragraph = doc.paragraphs[summary_idx]
    emphasis_run = next(run for run in summary_paragraph.runs if "優勢評估分數較高" in run.text)
    assert emphasis_run.bold is True
    assert emphasis_run.font.color.rgb is not None
    assert str(emphasis_run.font.color.rgb) == "ED0000"

    all_table_text = [cell.text for t in doc.tables for row in t.rows for cell in row.cells]
    assert "◆ 建議內容可依實際需求補充。" not in all_table_text


def test_apply_recommendation_format_overrides_sets_high_and_low_intro_line_spacing_to_16pt() -> None:
    from docx import Document
    from docx.shared import Pt

    converter = WordReportConverter()
    doc = Document()
    high_intro = (
        "感謝您接受心理潛能細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，"
        "空間感等共1項優勢評估分數較高，在此，也提供給您改善及建議方針："
    )
    low_intro = (
        "感謝您接受健康趨勢細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，"
        "空間感等共1項優勢評估分數較低，在此，也提供給您改善及建議方針："
    )
    high_paragraph = doc.add_paragraph(high_intro)
    low_paragraph = doc.add_paragraph(low_intro)

    converter._apply_recommendation_format_overrides(doc)

    assert high_paragraph.paragraph_format.line_spacing == Pt(16)
    assert low_paragraph.paragraph_format.line_spacing == Pt(16)


def test_apply_recommendation_format_overrides_sets_disclaimer_font_size_to_10pt() -> None:
    from docx import Document
    from docx.shared import Pt

    converter = WordReportConverter()
    doc = Document()
    disclaimer = (
        "本報告所提供之心理天賦優勢分析，係依據分子生物學資料及統計模型，推估個人心理天賦特質，"
        "僅供潛能探索與心理管理參考。"
    )
    paragraph = doc.add_paragraph(disclaimer)

    converter._apply_recommendation_format_overrides(doc)

    assert paragraph.runs
    assert all(run.font.size == Pt(10) for run in paragraph.runs)


def test_replace_recommendation_section_inserts_19pt_blank_paragraph_between_high_intro_and_first_high_table() -> None:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.shared import Pt

    converter = WordReportConverter()
    doc = Document()
    doc.add_paragraph("本報告所提供之心理天賦優勢分析")
    low_intro = "感謝您接受心理潛能細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，低特質等共1項優勢評估分數較低，在此，也提供給您改善及建議方針："
    doc.add_paragraph(low_intro)

    table = doc.add_table(rows=3, cols=6)
    headers = ["編號", "功能", "細胞解碼位點", "解碼型", "健康優勢評估", "健康優勢評分"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    table.rows[1].cells[1].text = "服從性格"
    table.rows[1].cells[4].text = "高"
    table.rows[2].cells[1].text = "低特質"
    table.rows[2].cells[4].text = "低"

    converter._replace_recommendation_section(doc, "王曉明")
    converter._apply_recommendation_format_overrides(doc)

    high_intro_paragraph = next(p for p in doc.paragraphs if "優勢評估分數較高，在此，也提供給您改善及建議方針：" in p.text)
    blank_paragraph_element = high_intro_paragraph._p.getnext()
    blank_paragraph = Paragraph(blank_paragraph_element, high_intro_paragraph._parent)

    assert blank_paragraph.text == ""
    assert blank_paragraph.paragraph_format.line_spacing == Pt(19)

    first_high_table_element = blank_paragraph_element.getnext()
    assert first_high_table_element.tag.endswith("tbl")
    first_high_table_text = "".join(first_high_table_element.itertext())
    assert "服從性格" in first_high_table_text


def test_replace_recommendation_section_keeps_two_greetings_high_then_low_and_page_break() -> None:
    converter = WordReportConverter()
    main_table = _build_main_table()
    main_table.rows[1].cells[1].text = "爆發力"
    main_table.rows[1].cells[4].text = "高"
    main_table.rows[2].cells[1].text = "低特質"
    main_table.rows[2].cells[4].text = "低"
    low_intro = (
        "感謝您接受健康趨勢細胞解碼檢測，由檢測結果得知，您在此次的分析項目中，"
        "低特質等共1項優勢評估分數較低，在此，也提供給您改善及建議方針："
    )
    doc = FakeDocument(
        paragraphs=[
            FakeParagraph("本報告所提供之心理天賦優勢分析"),
            FakeParagraph("_____舊名字_____ 貴賓您好："),
            FakeParagraph(low_intro),
        ],
        tables=[main_table],
    )

    converter._replace_recommendation_section(doc, "王曉明")

    texts = [p.text for p in doc.paragraphs]
    greeting = "_____王曉明_____ 貴賓您好："
    greeting_indexes = [idx for idx, text in enumerate(texts) if text == greeting]
    high_intro_index = next(idx for idx, text in enumerate(texts) if "優勢評估分數較高" in text)
    low_intro_index = next(idx for idx, text in enumerate(texts) if "優勢評估分數較低" in text)

    assert len(greeting_indexes) == 2
    assert greeting_indexes[0] < high_intro_index
    assert greeting_indexes[1] < low_intro_index
    assert any(text == "\f" for text in texts[greeting_indexes[0] : greeting_indexes[1]])


def test_convert_real_sample_docx_does_not_crash_and_keeps_two_greetings(tmp_path) -> None:
    import base64
    import shutil
    import re
    from docx import Document

    converter = WordReportConverter()
    sample_src = Path("src/word_converter/samples/input/APT-01-009297.docx")
    sample = tmp_path / sample_src.name
    shutil.copy(sample_src, sample)
    tiny_png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0r8AAAAASUVORK5CYII="
    )
    (tmp_path / "威力logo總表.png").write_bytes(tiny_png)
    (tmp_path / "心理logo總表.png").write_bytes(tiny_png)

    result = converter.convert(sample, tmp_path)

    assert result.output_path.exists()

    output_doc = Document(str(result.output_path))
    greeting_pattern = re.compile(r"^_+.+_+\s*貴賓您好：\s*$")
    greetings = [p.text.strip() for p in output_doc.paragraphs if greeting_pattern.match(p.text.strip())]
    assert len(greetings) == 2


def test_apply_first_page_logos_raises_when_no_body_image_paragraph(tmp_path) -> None:
    import base64
    from docx import Document

    converter = WordReportConverter()
    tiny_png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0r8AAAAASUVORK5CYII="
    )
    (tmp_path / "威力logo總表.png").write_bytes(tiny_png)
    (tmp_path / "心理logo總表.png").write_bytes(tiny_png)

    doc = Document()
    doc.add_paragraph("這份文件沒有圖片段落")

    with pytest.raises(ValueError, match="找不到正文第一個含圖片"):
        converter._apply_first_page_logos(doc, tmp_path)


def test_convert_inline_to_floating_anchor_works_without_cnvgraphicframepr_attribute() -> None:
    from types import SimpleNamespace

    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    converter = WordReportConverter()
    drawing = parse_xml(
        f"""
        <w:drawing {nsdecls('w', 'wp', 'a', 'pic', 'r')}>
          <wp:inline>
            <wp:extent cx="3685039" cy="1004563"/>
            <wp:docPr id="1" name="Picture 1"/>
            <a:graphic>
              <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic/>
              </a:graphicData>
            </a:graphic>
          </wp:inline>
        </w:drawing>
        """
    )
    inline = drawing[0]
    inline_shape = SimpleNamespace(_inline=inline)

    converter._convert_inline_to_floating_anchor(
        inline_shape,
        x_cm=8.76,
        y_cm=0,
        horizontal_relative="rightMargin",
        vertical_relative="paragraph",
        behind_text=True,
    )

    assert drawing[0].tag.endswith("anchor")
    assert drawing[0].find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}cNvGraphicFramePr") is not None
